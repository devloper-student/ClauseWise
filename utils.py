import io
from datetime import datetime
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_pdf(analysis_results: Dict[str, Any]) -> bytes:
    """
    Export analysis results to PDF format
    
    Args:
        analysis_results: Dictionary containing analysis data
        
    Returns:
        bytes: PDF file content
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4e79'),
        spaceAfter=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1f4e79'),
        spaceBefore=20,
        spaceAfter=10
    )
    
    # Build PDF content
    story = []
    
    # Title
    story.append(Paragraph("ClauseWise Analysis Report", title_style))
    story.append(Spacer(1, 20))
    
    # Document information
    story.append(Paragraph("Document Information", heading_style))
    doc_info = [
        ['Filename:', analysis_results.get('filename', 'N/A')],
        ['Analysis Date:', analysis_results.get('upload_time', datetime.now()).strftime('%Y-%m-%d %H:%M:%S')],
        ['Total Clauses:', str(len(analysis_results.get('clauses', [])))]
    ]
    
    doc_table = Table(doc_info, colWidths=[2*inch, 4*inch])
    doc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(doc_table)
    story.append(Spacer(1, 20))
    
    # Risk Summary
    if 'risk_analysis' in analysis_results:
        risk_data = analysis_results['risk_analysis']
        story.append(Paragraph("Risk Summary", heading_style))
        
        risk_summary = [
            ['Overall Risk Score:', f"{risk_data.get('overall_risk_score', 0):.1f}%"],
            ['Completeness Score:', f"{risk_data.get('completeness_score', 0):.1f}%"],
            ['High Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('high', 0))],
            ['Medium Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('medium', 0))],
            ['Low Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('low', 0))]
        ]
        
        risk_table = Table(risk_summary, colWidths=[2*inch, 4*inch])
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(risk_table)
        story.append(Spacer(1, 20))
    
    # Clauses
    story.append(Paragraph("Clause Analysis", heading_style))
    
    for i, clause in enumerate(analysis_results.get('clauses', [])[:10]):  # Limit to first 10 clauses
        story.append(Paragraph(f"Clause {clause.get('id', i+1)}: {clause.get('category', 'General')}", styles['Heading3']))
        story.append(Paragraph(f"<b>Risk Level:</b> {clause.get('risk_level', 'Unknown').title()}", styles['Normal']))
        story.append(Paragraph(f"<b>Summary:</b> {clause.get('simplified_text', 'No summary available')}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_word(analysis_results: Dict[str, Any]) -> bytes:
    """
    Export analysis results to Word format
    
    Args:
        analysis_results: Dictionary containing analysis data
        
    Returns:
        bytes: Word document content
    """
    document = DocxDocument()
    
    # Add title
    title = document.add_heading('ClauseWise Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document information
    document.add_heading('Document Information', level=1)
    doc_info = document.add_table(rows=3, cols=2)
    doc_info.style = 'Table Grid'
    
    cells = doc_info.rows[0].cells
    cells[0].text = 'Filename:'
    cells[1].text = analysis_results.get('filename', 'N/A')
    
    cells = doc_info.rows[1].cells
    cells[0].text = 'Analysis Date:'
    cells[1].text = analysis_results.get('upload_time', datetime.now()).strftime('%Y-%m-%d %H:%M:%S')
    
    cells = doc_info.rows[2].cells
    cells[0].text = 'Total Clauses:'
    cells[1].text = str(len(analysis_results.get('clauses', [])))
    
    # Risk Summary
    if 'risk_analysis' in analysis_results:
        risk_data = analysis_results['risk_analysis']
        document.add_heading('Risk Summary', level=1)
        
        risk_table = document.add_table(rows=5, cols=2)
        risk_table.style = 'Table Grid'
        
        risk_rows = [
            ('Overall Risk Score:', f"{risk_data.get('overall_risk_score', 0):.1f}%"),
            ('Completeness Score:', f"{risk_data.get('completeness_score', 0):.1f}%"),
            ('High Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('high', 0))),
            ('Medium Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('medium', 0))),
            ('Low Risk Clauses:', str(risk_data.get('risk_breakdown', {}).get('low', 0)))
        ]
        
        for i, (label, value) in enumerate(risk_rows):
            cells = risk_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
    
    # Recommendations
    if 'risk_analysis' in analysis_results and 'recommendations' in analysis_results['risk_analysis']:
        document.add_heading('Recommendations', level=1)
        for recommendation in analysis_results['risk_analysis']['recommendations'][:10]:
            p = document.add_paragraph(recommendation, style='List Bullet')
    
    # Clauses
    document.add_heading('Clause Analysis', level=1)
    
    for clause in analysis_results.get('clauses', [])[:15]:  # Limit to first 15 clauses
        document.add_heading(f"Clause {clause.get('id', '')}: {clause.get('category', 'General')}", level=2)
        
        # Risk level
        risk_paragraph = document.add_paragraph()
        risk_paragraph.add_run('Risk Level: ').bold = True
        risk_paragraph.add_run(clause.get('risk_level', 'Unknown').title())
        
        # Summary
        summary_paragraph = document.add_paragraph()
        summary_paragraph.add_run('Summary: ').bold = True
        summary_paragraph.add_run(clause.get('simplified_text', 'No summary available'))
        
        # Original text (truncated)
        original_text = clause.get('text', '')
        if len(original_text) > 500:
            original_text = original_text[:500] + "..."
        
        original_paragraph = document.add_paragraph()
        original_paragraph.add_run('Original Text: ').bold = True
        original_paragraph.add_run(original_text)
        
        document.add_paragraph()  # Add space
    
    # Save to buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def format_risk_level(risk_level: str) -> str:
    """Format risk level with emoji"""
    emojis = {
        'high': '🔴',
        'medium': '🟡',
        'low': '🟢'
    }
    return f"{emojis.get(risk_level.lower(), '⚪')} {risk_level.title()}"

def truncate_text(text: str, max_length: int = 200) -> str:
    """Truncate text to specified length"""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."

def get_risk_color_hex(risk_level: str) -> str:
    """Get hex color for risk level"""
    colors = {
        'high': '#ff4444',
        'medium': '#ffaa00',
        'low': '#00aa44'
    }
    return colors.get(risk_level.lower(), '#888888')
