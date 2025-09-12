import io
import pdfplumber
from docx import Document
import streamlit as st
from typing import Optional

def parse_document(uploaded_file) -> str:
    """
    Parse text content from uploaded PDF, DOCX, or TXT files
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    try:
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return parse_pdf(uploaded_file)
        elif file_extension == 'docx':
            return parse_docx(uploaded_file)
        elif file_extension == 'txt':
            return parse_txt(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        st.error(f"Error parsing document: {str(e)}")
        return ""

def parse_pdf(uploaded_file) -> str:
    """Extract text from PDF file using pdfplumber"""
    try:
        pdf_bytes = uploaded_file.read()
        text_content = ""
        
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text + "\n"
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
        
        return text_content.strip()
        
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")

def parse_docx(uploaded_file) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        docx_bytes = uploaded_file.read()
        text_content = ""
        
        doc = Document(io.BytesIO(docx_bytes))
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
                text_content += "\n"
        
        return text_content.strip()
        
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")

def parse_txt(uploaded_file) -> str:
    """Extract text from TXT file"""
    try:
        text_content = uploaded_file.read()
        
        # Try to decode as UTF-8 first, then fall back to other encodings
        if isinstance(text_content, bytes):
            try:
                text_content = text_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = text_content.decode('latin-1')
                except UnicodeDecodeError:
                    text_content = text_content.decode('utf-8', errors='ignore')
        
        return text_content.strip()
        
    except Exception as e:
        raise Exception(f"Error parsing TXT: {str(e)}")

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            cleaned_lines.append(line)
    
    # Join lines and normalize spacing
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Remove multiple consecutive newlines
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text

def validate_document_content(text: str) -> bool:
    """Validate that the extracted text appears to be a legal document"""
    if not text or len(text.strip()) < 100:
        return False
    
    # Check for common legal terms (basic validation)
    legal_indicators = [
        'agreement', 'contract', 'party', 'parties', 'terms', 'conditions',
        'liability', 'indemnity', 'confidential', 'termination', 'clause',
        'section', 'whereas', 'therefore', 'herein', 'hereby'
    ]
    
    text_lower = text.lower()
    legal_term_count = sum(1 for term in legal_indicators if term in text_lower)
    
    # Document should contain at least 3 legal terms to be considered valid
    return legal_term_count >= 3
